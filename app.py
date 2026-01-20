from flask import Flask, request, jsonify, session
from werkzeug.utils import secure_filename
from extensions import db
from models import *
import os
import re
import pdfplumber
import docx
import json
import requests
from flask_cors import CORS
import fitz
from flask_migrate import Migrate
from collections import Counter

DEMO_MODE = True 

app = Flask(__name__)
# Allow ALL origins
CORS(app, resources={r"/*": {"origins": "*"}})

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# CONFIG
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///site.db"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = "super_secret_key_for_hackathon"
app.config.update(
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=False, 
    SESSION_COOKIE_HTTPONLY=True, 
    SESSION_PERMANENT=True
)

db.init_app(app)
migrate = Migrate(app, db)

with app.app_context():
    db.create_all()

# --- HELPER FUNCTIONS ---

def get_or_create_demo_user():
    demo_email = "demo@judge.com"
    user = Users.query.filter_by(email=demo_email).first()
    if not user:
        user = Users(email=demo_email, password_hash="demo", role="candidate")
        db.session.add(user)
        db.session.commit()
        candidate = Candidates(user_id=user.user_id, full_name="Demo Judge", experience_years=10)
        db.session.add(candidate)
        db.session.commit()
    return user

def extract_text_and_links(file_path):
    text = ""
    links = []
    try:
        if file_path.endswith(".pdf"):
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            doc = fitz.open(file_path)
            for page in doc:
                for link in page.get_links():
                    if link.get("uri"): links.append(link.get("uri"))
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for rel in doc.part.rels.values():
                if "http" in rel.target_ref:
                    links.append(rel.target_ref)
    except Exception as e:
        print(f"Error reading file: {e}")
    return text.lower(), links

KNOWN_SKILLS = ["python", "c", "c++", "java", "javascript", "html", "css", "sql", "flask", "django", "react", "node", "machine learning", "deep learning", "go", "golang", "rust", "kubernetes", "docker", "aws", "azure"]

def extract_skills(text):
    found_skills = set()
    for skill in KNOWN_SKILLS:
        if re.search(r"\b" + re.escape(skill) + r"\b", text):
            found_skills.add(skill.capitalize())
    return list(found_skills)

def extract_github_username(text, links):
    match = re.search(r"github\.com/([a-zA-Z0-9_-]+)", text)
    if match: return match.group(1)
    for link in links:
        match = re.search(r"github\.com/([a-zA-Z0-9_-]+)", link)
        if match: return match.group(1)
    return None

def extract_email(text):
    match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    return match.group(0) if match else "Not found"

# Simple cleaning
def clean_extracted_text(text):
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_education(text):
    keywords = ["university", "college", "institute", "b.tech", "b.sc", "degree", "bachelor", "master", "phd"]
    # Simple line-based extraction
    lines = text.split('\n')
    education_lines = []
    for line in lines:
        if any(word in line.lower() for word in keywords):
            clean = clean_extracted_text(line)
            if len(clean) > 10 and len(clean) < 150:
                education_lines.append(clean)
    return education_lines[:2]

def extract_section(text, header_keywords):
    # A simple mock extractor
    lines = text.split('\n')
    captured_lines = []
    capture = False
    for line in lines:
        clean_line = line.strip().lower()
        if any(w in clean_line for w in header_keywords) and len(clean_line) < 40:
            capture = True
            continue
        if capture:
            # Stop if we hit another header
            if any(w in clean_line for w in ["education", "skills", "projects", "experience", "certifications"]):
                break
            if len(line.strip()) > 3:
                captured_lines.append(clean_extracted_text(line))
    return captured_lines[:5]

def extract_keywords(text):
    # Count most frequent meaningful words (simple "tag cloud")
    words = re.findall(r'\b[a-zA-Z]{5,}\b', text.lower())
    ignore = {"experience", "project", "technologies", "worked", "using", "months", "years", "university"}
    filtered = [w for w in words if w not in ignore]
    return [item[0] for item in Counter(filtered).most_common(8)]

def parse_resume(file_path):
    text_lower, links = extract_text_and_links(file_path)
    
    # Get original casing for display
    raw_text_original = ""
    try:
        if file_path.endswith(".pdf"):
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    raw_text_original += page.extract_text() or ""
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                raw_text_original += para.text + "\n"
    except:
        raw_text_original = text_lower

    if not raw_text_original: raw_text_original = text_lower

    return {
        "skills": extract_skills(text_lower),
        "user": extract_github_username(text_lower, links),
        "email": extract_email(text_lower),
        "phone": "Not found", # Simplify for demo
        "education": extract_education(raw_text_original),
        "experience": extract_section(raw_text_original, ["experience", "work history"]),
        "projects": extract_section(raw_text_original, ["projects"]),
        "raw_text": raw_text_original,
        "keywords": extract_keywords(raw_text_original)  # <--- NEW FEATURE
    }

# --- ROUTES ---

@app.route("/dashboard")
def dashboard():
    if DEMO_MODE and "user_id" not in session:
        user = get_or_create_demo_user()
        session["user_id"] = user.user_id
        session["role"] = user.role
    if "user_id" not in session: return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"message": "Welcome to Dashboard"})

@app.route("/parse_resume", methods=["POST"])
def parse_resume_endpoint():
    if DEMO_MODE and "user_id" not in session:
        user = get_or_create_demo_user()
        session["user_id"] = user.user_id

    resume_file = request.files.get("resume")
    if not resume_file: return jsonify({"error": "No resume uploaded"}), 400

    filename = secure_filename(resume_file.filename)
    path = os.path.join(UPLOAD_FOLDER, filename)
    resume_file.save(path)

    parsed_data = parse_resume(path)
    
    return jsonify({
        "message": "Parsed successfully",
        "profile": {
            "user": parsed_data["user"] or "Vishalfot",
            "email": parsed_data["email"],
            "phone": parsed_data["phone"],
            "education": parsed_data["education"],
            "experience": parsed_data["experience"],
            "projects": parsed_data["projects"],
            "skills_found": parsed_data["skills"],
            "raw_text": parsed_data["raw_text"], # <--- Correctly placed here
            "keywords": parsed_data["keywords"]  # <--- New Feature
        }
    })

@app.route("/analyze_skills", methods=["POST"])


if __name__ == "__main__":
    app.run(debug=True, port=5000)